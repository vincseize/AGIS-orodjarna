# -*- coding: utf-8 -*-

"""
***************************************************************************
*                                                                         *
*   This program is free software; you can redistribute it and/or modify  *
*   it under the terms of the GNU General Public License as published by  *
*   the Free Software Foundation; either version 2 of the License, or     *
*   (at your option) any later version.                                   *
*                                                                         *
***************************************************************************
Matjaž Mori, ZVKDS CPA
19.2.2020

"""

from PyQt5.QtCore import QCoreApplication, QObject,QFileInfo
from PyQt5.QtGui import QIcon
from qgis.utils import iface
from qgis.core import (Qgis,
                        QgsProcessing,
                       QgsFeatureSink,
                       QgsMessageLog,
                       QgsProcessingException,
                       QgsProcessingAlgorithm,
                       QgsProcessingParameterFolderDestination,
                       QgsProcessingParameterFeatureSource,
                       QgsProcessingParameterFeatureSink,
                       QgsApplication,
                       QgsProcessingParameterEnum,
                       QgsProcessingParameterBoolean,
                       NULL)
import processing
import datetime
import sys
import os

from ..externals import path, checkDuplicates, value_error
from pathlib import Path


try:
    import docx
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import RGBColor
    from docx.shared import Length
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    iface.messageBar().pushMessage("DOCX module found!", duration=3)
except:
    import subprocess
    subprocess.check_call(['python', '-m', 'pip', 'install', 'python-docx'])
    iface.messageBar().pushMessage("Error", "No docx module found, installing!", level=Qgis.Critical, duration=10)


class StratiWordCatalog(QgsProcessingAlgorithm):
    """
    This is algorithm that takes a vector layer and
    creates catalog in docx file.
    """

    # Constants used to refer to parameters and outputs. They will be
    # used when calling the algorithm from another algorithm, or when
    # calling from the QGIS console.

    KATALOGSE = 'KATALOGSE'
    FOLDER = 'FOLDER'
    OUTPUT = 'OUTPUT'
    SORT_BY_FAZE = 'SORT_BY_FAZE'
    POVZETEK = 'POVZETEK'
    OPCIJE = 'OPCIJE'

    def tr(self, string):
        """
        Returns a translatable string with the self.tr() function.
        """
        return QCoreApplication.translate('Processing', string)

    def createInstance(self):
        return StratiWordCatalog()

    def name(self):
        """
        Returns the algorithm name, used for identifying the algorithm. This
        string should be fixed for the algorithm, and must not be localised.
        The name should be unique within each provider. Names should contain
        lowercase alphanumeric characters only and no spaces or other
        formatting characters.
        """
        return 'se_catalog'

    def displayName(self):
        """
        Returns the translated algorithm name, which should be used for any
        user-visible display of the algorithm name.
        """
        return self.tr('Seznam SE > docx')

    def group(self):
        """
        Returns the name of the group this algorithm belongs to. This string
        should be localised.
        """
        return self.tr(u'Poročilo')
        pass

    def groupId(self):
        """
        Returns the unique ID of the group this algorithm belongs to. This
        string should be fixed for the algorithm, and must not be localised.
        The group id should be unique within each provider. Group id should
        contain lowercase alphanumeric characters only and no spaces or other
        formatting characters.
        """
        return 'porocilo'
        pass

    def shortHelpString(self):
        """
        Returns a localised short helper string for the algorithm. This string
        should provide a basic description about what the algorithm does and the
        parameters and outputs associated with it..
        """

        help = ("""
            To orodje pretvori seznam z opisi stratigrafskih enot v kataloški izpis v wordovo datoteko.
            <h3></h3>
            <ul>
            </ul>
            """)
        return self.tr(help)

    def icon(self):
        """
        Should return a QIcon which is used for your provider inside
        the Processing toolbox.
        """
        return QIcon(os.path.join(path('icons'),'strati_word_catalog.png'))

    def initAlgorithm(self, config=None):
        """
        Here we define the inputs and output of the algorithm, along
        with some other properties.
        """

        # We add the input vector features source. It can have any kind of
        # geometry.
        self.addParameter(
            QgsProcessingParameterFeatureSource(
                self.KATALOGSE,
                self.tr('Opisi SE'),
                [QgsProcessing.TypeVector]
            )
        )

        self.addParameter(
            QgsProcessingParameterFolderDestination(
                self.FOLDER,
                self.tr('Ciljna mapa kataloga'),
                defaultValue= 'C:/Users/matjaz.mori.ZVKDS/Desktop'
                )
        )

        self.addParameter(
            QgsProcessingParameterBoolean(
                self.SORT_BY_FAZE,
                self.tr('Sortiraj po Fazah'),
                defaultValue=True
                )
        )

        self.addParameter(
            QgsProcessingParameterBoolean(
                self.POVZETEK,
                self.tr('Vključi povzetek po Fazah'),
                defaultValue=True
                )
        )


        self.addParameter(
            QgsProcessingParameterEnum(
                self.OPCIJE,
                'Dodatne možnosti',
                optional=True,
                options=['Vključi datacijo','Vključi kvadrante', 'Vključi sektor', 'Vključi opažanja'],
                allowMultiple=True,
                defaultValue=None
                )
        )


    def processAlgorithm(self, parameters, context, feedback):
        """
        Here is where the processing itself takes place.
        """

        # Retrieve the feature source and sink. The 'dest_id' variable is used
        # to uniquely identify the feature sink, and must be included in the
        # dictionary returned by the processAlgorithm function.
        source = self.parameterAsSource(
            parameters,
            self.KATALOGSE,
            context
        )

        output = self.parameterAsString(
            parameters,
            self.FOLDER,
            context
        )

        sort_by_faze = self.parameterAsBool(
            parameters,
            self.SORT_BY_FAZE,
            context
        )

        povzetek_by_faze = self.parameterAsBool(
            parameters,
            self.POVZETEK,
            context
        )

        # Prepeard to upgrade with options

        opcije = self.parameterAsEnums(
            parameters,
            self.OPCIJE,
            context
        )


        # If source was not found, throw an exception to indicate that the algorithm
        # encountered a fatal error. The exception text can be any string, but in this
        # case we use the pre-built invalidSourceError method to return a standard
        # helper text for when a source cannot be evaluated
        if source is None:
            raise QgsProcessingException(self.invalidSourceError(parameters, self.KATALOG))

        # Send some information to the user
        feedback.pushInfo(output)


        # Compute the number of steps to display within the progress bar and
        # get features from source
        total = 100.0 / source.featureCount() if source.featureCount() else 0
        features = source.getFeatures()


        #Check duplicates of SE
        checkDuplicates(features, 'se_id', feedback)

        #Set output file name
        now = datetime.datetime.now()
        date = now.strftime("%Y%m%d")
        docx = '\Katalog stratigrafskih enot_' + date + '.docx'

        # python docx
        document = Document()
        dockatalog = output + docx

        # Define text style
        style = document.styles['Normal']
        font = style.font
        font.name = 'Georgia'
        font.size = Pt(10)
        style.hidden = False
        style.quick_style = True
        style.priority = 1

        #Define title font
        styles = document.styles
        new_heading_style = styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = styles['Heading 4']
        fonth = new_heading_style.font
        fonth.name = 'Georgia'
        fonth.size = Pt(10)
        fonth.color.rgb = RGBColor(0x0, 0x0, 0x0)
        new_heading_style.hidden = False
        new_heading_style.quick_style = True
        new_heading_style.priority = 1

        #Define sub-title font
        styles = document.styles
        sub_heading_style = styles.add_style('Sub Heading', WD_STYLE_TYPE.PARAGRAPH)
        sub_heading_style.base_style = styles['Heading 5']
        fonth = sub_heading_style.font
        fonth.name = 'Georgia'
        fonth.size = Pt(10)
        fonth.color.rgb = RGBColor(0x0, 0x0, 0x0)
        sub_heading_style.hidden = False
        sub_heading_style.quick_style = True
        sub_heading_style.priority = 1

        #Get unique values of "faza" values
        idx = source.fields().indexOf('faza')
        faze = source.uniqueValues(idx)
        if NULL in faze and '' in faze:
            faze.remove('')





        # Define title text
        def naslov_id(field):
            if field is '' or field == NULL:
                field = 'ni določena!'
                return field
            else:
                return field

        #Write SE

        def print_se(feature):

            se_id = feature['se_id']
            vrsta = feature['vrsta']
            dolocljivost_meje = feature['dolocljivost_meje']
            oblika_tloris = feature['oblika_tloris']
            sektor = feature['sektor']
            kvadrant = feature['kvadrant']
            vrsta_posega = feature['vrsta_posega']
            oblika_profil = feature['obllika_profil']
            dolzina = feature['dolzina']
            sirina = feature['sirina']
            debelina = feature['debelina']

            barva = feature['barva']
            barva_munsel = feature['barva_munsel']
            konsistenca = feature['konsistenca']
            tekstura = feature['tekstura']
            grobe_sestavine  = feature['grobe_sestavine']
            najdbe = feature['najdbe']
            opis = feature['opis']
            nastanek = feature['nastanek']
            interpretacija = feature['interpretacija']
            faza = feature['faza']
            odstranitev = feature['odstranitev_se']
            datacija = feature['datacija']

            
           
            #define body parts
            def p_se_id(p):
                if se_id != None:
                    p.add_run('SE ').bold = True
                    p.add_run(str(naslov_id(se_id))).bold = True
                    p.add_run(', ')
                else:
                    value_error(se_id, 'se_id', feedback)

            def p_vrsta(p):
                if vrsta != None:
                    p.add_run(str(vrsta))
                    p.add_run(" ")
                elif "IZNIČENO" in str(feature['opis']):
                    pass
                else:
                    value_error(se_id, 'vrsta', feedback)

            def p_dolocljivost_meje(p):
                if dolocljivost_meje != None:
                    p.add_run("; definicija meje ")
                    p.add_run(str(dolocljivost_meje))
                else:
                    feedback.reportError('Pri SE %s manjaka vrednost dolocljivost_meje!' % se_id, False)
                    pass
                      
            def p_oblika_tloris(p):
                if oblika_tloris == oblika_profil and oblika_tloris != None:
                    p.add_run(str(oblika_tloris))
                    p.add_run(" oblika v tlorisu in preseku")
                if oblika_profil != oblika_tloris and oblika_tloris != None:          
                    p.add_run(str(oblika_tloris))
                    p.add_run(" oblika v tlorisu")
                else:
                    feedback.reportError('Pri SE %s manjaka vrednost oblika_tloris!' % se_id, False)
                    pass

            def p_oblika_profil(p):
                if oblika_profil != oblika_tloris and oblika_profil != None:      
                    if oblika_tloris != None:
                        p.add_run(" in ")
                    p.add_run(str(oblika_profil))
                    p.add_run(" oblika v preseku")
                else:
                    feedback.pushDebugInfo('Pri SE %s manjka oblika v preseku!' % se_id)
                    pass


            def dimensions(dimension):
                string = str(dimension)
                replaced = string.replace('.', ',')
                p.add_run(replaced)

            def p_velikost(p):
                if (sirina == None or sirina == 'NULL') and (dolzina == None or dolzina == 'NULL'):
                    feedback.pushDebugInfo('Pri SE %s manjkajo dimenzije!' % se_id)
                    pass

                elif (sirina == None or sirina == 'NULL') and (oblika_tloris != 'ovalna' or oblika_tloris != 'okrogla'):
                    p.add_run(", premera ")
                    dimensions(dolzina)
                    p.add_run(" m")

                elif (dolzina == None or dolzina == 'NULL') and (oblika_tloris != 'ovalna' or oblika_tloris != 'okrogla'):
                    p.add_run(", premera ")
                    dimensions(sirina)
                    p.add_run(" m")
        
                else:
                    p.add_run(", velikosti ")
                    dimensions(dolzina)
                    p.add_run(" x ")
                    dimensions(sirina)
                    p.add_run(" m")

            def p_debelina(p):
                if debelina != None and vrsta in ('plast', 'mejna površina', 'polnilo') :
                    p.add_run(", debeline do ")
                    dimensions(debelina)
                    p.add_run(" m")

                elif debelina != None and vrsta == 'vkop' :
                    p.add_run(", globine do ")
                    dimensions(debelina)
                    p.add_run(" m")    

                elif debelina != None and vrsta == 'struktura' :
                    p.add_run(", višine do ")
                    dimensions(debelina)
                    p.add_run(" m")    

                elif debelina != None:
                    p.add_run(", debeline do ")
                    dimensions(debelina)
                    p.add_run(" m")

                else:
                    feedback.pushDebugInfo('Pri SE %s manjka debelina!' % se_id)
                    pass


            def p_konsistenca(p):
                if konsistenca != None:
                    #p.add_run("; konsistenca: ")
                    p.add_run(str(konsistenca))
                    p.add_run(" ")
                else:
                    pass

            def p_tekstura(p):
                if tekstura != None:
                    #p.add_run("; tekstura: ")
                    p.add_run(str(tekstura))
                    p.add_run(" ")
                else:
                    pass

            def p_barva(p):
                if barva != None:
                    #p.add_run("; barva: ")
                    p.add_run(str(barva))
                    if barva_munsel != None:
                        p.add_run(' (')
                        p.add_run(str(barva_munsel))
                        p.add_run('), ')
                elif barva == None and vrsta in ('plast', 'polnilo'):
                    feedback.reportError('Pri SE %s manjaka barva!' % se_id, False)
                    pass
                else:
                    pass

            def p_grobe_sestavine(p):
                if grobe_sestavine != None or grobe_sestavine != NULL:
                    p.add_run("z/s ")
                    sest = grobe_sestavine.replace('\n', ', ')
                    p.add_run(str(sest))
                    p.add_run(" ")
                else:
                    pass

            def p_opis(p):
                if opis != None:
                    p.add_run("Opažanja: ")
                    p.add_run(str(opis))
                else:
                    pass

            def p_najdbe(p):
                p.add_run("Stratigrafska enota ")
                if najdbe:
                    p.add_run("je vsebovala najdbe")
                else:
                    p.add_run("ni vsebovala najdb")

            def p_odstranitev(p):
                if odstranitev == 'ohranjena':
                    p.add_run(" in ni bila odstranjena.")
                elif odstranitev == 'delno odstranjena':
                    poseg = p_vrsta_posega()
                    p.add_run(" in je bila %s delno odstranjena." % poseg)
                elif odstranitev == 'odstranjena':
                    poseg = p_vrsta_posega()
                    p.add_run(" in je bila %s v celoti odstranjena." % poseg)
                elif odstranitev == NULL or odstranitev == None:
                    p.add_run(". ")

            def p_vrsta_posega():
                if vrsta_posega == 'strojno':
                    poseg = 'strojno'
                    return 'strojno'
                elif vrsta_posega == 'ročno/strojno':
                    poseg ='ročno/strojno'
                    return poseg
                elif vrsta_posega == 'ročno':
                    poseg ='ročno'
                    return poseg
                else:
                    poseg =''
                    return poseg
                    pass

            def p_faza(p):
                p.add_run("Faza: ")
                if faza != None and faza != NULL:
                    p.add_run(faza)
                    p.add_run(". ")
                else:
                    p.add_run("ni določena. ")

            def p_interpretacija(p):
                if interpretacija != None:
                    p.add_run("Interpretacija: ")
                    p.add_run(interpretacija)
                    p.add_run(".\n")
                else:
                    p.add_run("Interpretacija: ")
                    p.add_run("-.\n")

            def p_datacija(p):
                if datacija != None:
                    p.add_run("Datacija: ")
                    p.add_run(datacija)
                    p.add_run(". ")
                else:
                    pass

            def p_kvadranti(p):
                if kvadrant != None:
                    p.add_run("Kv.: ")
                    p.add_run(kvadrant)
                    p.add_run("; ")
                else:
                    pass

            def p_sektor(p):
                if sektor != None:
                    p.add_run("Sektor: ")
                    p.add_run(sektor)
                    p.add_run("; ")
                else:
                    pass

            def p_enaka(p):
                if sektor != None:
                    p.add_run("Enačena z ")
                    p.add_run(enaka)
                    p.add_run(". ")
                else:
                    pass


            #Define paragraphs

            p = document.add_paragraph(style = 'Sub Heading')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format = p.paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.space_before = Pt(0)
            paragraph_format.line_spacing = 1.5
            body_text = p.style.name

            if 3 in opcije:
                po = document.add_paragraph(style = 'Normal')
                po.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
                paragraph_formatl = po.paragraph_format
                paragraph_formatl.space_after = Pt(0)
                paragraph_formatl.space_before = Pt(0)
                paragraph_formatl.line_spacing = 1.5
            """
            pn = document.add_paragraph()
            pn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
            paragraph_formatv = pn.paragraph_format
            paragraph_formatv.space_after = Pt(0)
            paragraph_formatv.space_before = Pt(0)
            paragraph_formatv.line_spacing = 1.5
            """
    
            pi = document.add_paragraph()
            pi.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_formatv = pi.paragraph_format
            paragraph_formatv.space_after = Pt(0)
            paragraph_formatv.space_before = Pt(0)
            paragraph_formatv.line_spacing = 1.5

            #Print Body, data
            p_se_id(p)
            if 2 in opcije:
                p_sektor(p)
            if 1 in opcije:
                p_kvadranti(p)
            
            #p_dolocljivost_meje(p)

            p_barva(p)
            p_konsistenca(p)
            p_vrsta(p)
            p_tekstura(p)
            p_grobe_sestavine(p)

            p_oblika_tloris(p)
            p_oblika_profil(p)
        
            p_velikost(p)
            p_debelina(p)

            
            p.add_run(".")



            #Print Body, description
            if 3 in opcije:
                p_opis(po)

            #p_najdbe(pn)
            #p_odstranitev(pn)

            if 0 in opcije:
                p_datacija(pi)
            """"
            if not sort_by_faze:
                p_faza(pi)
            """
            p_interpretacija(pi)
            

        def povzetek_faze():
            if povzetek_by_faze:
                p = document.add_paragraph(style = 'Heading')
                p.add_run("Povzetek SE po fazah")
                for faza in faze:
                    po = document.add_paragraph(style = 'Normal')
                    po.add_run("Faza ")
                    po.add_run(faza)
                    list_se = []
                    for current, feature in enumerate(source.getFeatures()):
                            if "IZNIČENO" in str(feature['opis']):
                                pass
                            else:
                                sum_text = "SE %s, %s" %(feature['se_id'], feature['interpretacija'])
                                list_se.append(sum_text)

                    p = document.add_paragraph(style = 'Normal')
                    p.add_run('; '.join(list_se))
                p = document.add_paragraph(style = 'Normal')
                p = document.add_paragraph(style = 'Normal')

        def print_each_se():
            povzetek_faze()
            for current, feature in enumerate(source.getFeatures()):
                feedback.setProgress(int(current * total))
                feedback.pushInfo('Zapisujem SE %s' % feature['se_id'])
                if "IZNIČENO" in str(feature['opis']):
                    feedback.reportError('SE %s je bila izničena!' % feature['se_id'], False)
                    pass
                else:
                    print_se(feature)

        def print_faze():
            povzetek_faze()
            for faza in faze:

                # Stop the algorithm if cancel button has been clicked
                if feedback.isCanceled():
                    break

                feedback.pushInfo('Zapisujem fazo %s' %faza)
                #Write title
                naslov = "Faza " + str(naslov_id(faza))
                n = document.add_paragraph(naslov, style='Heading')
                n.alignment = WD_ALIGN_PARAGRAPH.LEFT
                n.paragraph_format.space_before = Pt(10)
                n.paragraph_format.space_after = Pt(10)
                n.paragraph_format.line_spacing = 1.5

                #Write each SE if matches "faza"
                for current, feature in enumerate(source.getFeatures()):
                    feedback.setProgress(int(current * total))
                    if feature['faza'] == faza:
                        if "IZNIČENO" in str(feature['opis']):
                            feedback.reportError('SE %s je bila izničena!' % feature['se_id'], False)
                            pass
                        else:
                            print_se(feature)

        if sort_by_faze:
            print_faze()
        else:
            print_each_se()





        # Return the results of the algorithm. In this case our only result is
        # the feature sink which contains the processed features, but some
        # algorithms may return multiple feature sinks, calculated numeric
        # statistics, etc. These should all be included in the returned
        # dictionary, with keys matching the feature corresponding parameter
        # or output names.
        feedback.pushInfo(dockatalog)
        document.save(dockatalog)

        results ={}

        return results
